"""Machine document ingestion — upload a brochure, get a searchable machine.

This is what makes the catalog the client's to own. Without it, adding a
machine means editing a markdown file in the repo and running a script from a
developer's laptop, which is not a product anyone can operate.

Flow: upload -> extract text -> store in machine_documents -> chunk -> embed ->
upsert into Qdrant. The machine is answerable in chat immediately afterwards.

Memory matters here: Render's free tier has 512MB and a large PDF can exhaust
it. Pages are extracted one at a time and embeddings run in small batches
rather than loading everything at once.
"""

from __future__ import annotations

import base64
import io
import re
from typing import Any

from pydantic import BaseModel, Field, ValidationError, create_model

from app import store
from app.config import settings
from app.enums import DocumentType
from app.llm import LLMUnavailableError, complete, embed, transcribe_image
from app.logging_config import get_logger
from app.rag import ensure_collection, extract_codes, get_client

log = get_logger(__name__)

# Beyond this the free tier risks an out-of-memory restart mid-upload.
MAX_UPLOAD_BYTES = 10 * 1024 * 1024

# Roughly 300 words. Long enough to hold a complete spec, short enough that a
# retrieved chunk is mostly relevant.
CHUNK_CHARS = 1800
CHUNK_OVERLAP = 200

SUPPORTED = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "text/markdown": "txt",
}


class ExtractionError(RuntimeError):
    """The file could not be turned into text. The message reaches the user."""


async def extract_text(data: bytes, filename: str, content_type: str | None) -> str:
    """Pull plain text out of an uploaded document.

    Raises ExtractionError with a message worth showing a user — "this looks
    like a scanned PDF" is actionable, "extraction failed" is not.
    """
    if len(data) > MAX_UPLOAD_BYTES:
        raise ExtractionError(
            f"File is {len(data) // 1024 // 1024}MB. The limit is "
            f"{MAX_UPLOAD_BYTES // 1024 // 1024}MB."
        )

    kind = SUPPORTED.get(content_type or "")
    if kind is None:
        lower = filename.lower()
        if lower.endswith(".pdf"):
            kind = "pdf"
        elif lower.endswith(".docx"):
            kind = "docx"
        elif lower.endswith((".txt", ".md")):
            kind = "txt"
        else:
            raise ExtractionError(
                "Unsupported file type. Upload a PDF, Word document, or text file."
            )

    if kind == "pdf":
        return await _extract_pdf(data)
    if kind == "docx":
        return _extract_docx(data)
    return data.decode("utf-8", errors="replace").strip()


async def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"Could not read the PDF: {exc}") from exc

    page_count = len(reader.pages)

    # Page by page, so a large document never sits in memory whole. A page
    # with no extractable text is usually a scan (a picture of a page, not
    # real text) rather than an empty page — those get OCR'd individually
    # rather than failing the whole document, since a brochure is often a
    # mix (a photo cover page, then real text pages, or vice versa).
    pages: list[str] = []
    ocr_pages: list[int] = []
    for index, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            pages.append(text.strip())
        else:
            pages.append("")
            ocr_pages.append(index)

    if ocr_pages:
        if page_count > settings.ocr_max_pages:
            log.warning(
                "ocr_skipped_too_many_pages",
                extra={"page_count": page_count, "limit": settings.ocr_max_pages},
            )
        else:
            ocr_results = await _ocr_pages(data, ocr_pages)
            for index, text in ocr_results.items():
                if text:
                    pages[index] = text

    combined = "\n\n".join(p for p in pages if p).strip()
    if not combined:
        # Every page failed both real extraction and OCR — genuinely nothing
        # to work with, rather than silently storing an empty document that
        # will never answer a question.
        raise ExtractionError(
            "No text could be read from this PDF, even with OCR. The scan "
            "quality may be too low, or the page is a diagram/photo with no "
            "text. Please upload a clearer PDF, or paste the specifications "
            "directly."
        )
    return combined


async def _ocr_pages(data: bytes, page_indices: list[int]) -> dict[int, str | None]:
    """Render specific pages to images and transcribe them via GPT-4o vision.

    Rendered one page at a time (not the whole document at once) to keep
    memory bounded — the same reasoning as the page-by-page text extraction
    above, just for images instead of text.
    """
    import pypdfium2 as pdfium

    results: dict[int, str | None] = {}
    try:
        pdf = pdfium.PdfDocument(data)
    except Exception:
        log.exception("ocr_pdf_open_failed")
        return results

    try:
        for index in page_indices:
            try:
                page = pdf[index]
                # 2x scale (~144 DPI) balances legibility for spec-sheet text
                # against image payload size sent to the vision API.
                bitmap = page.render(scale=2.0)
                pil_image = bitmap.to_pil()
                buffer = io.BytesIO()
                pil_image.save(buffer, format="PNG")
                image_b64 = base64.b64encode(buffer.getvalue()).decode()
            except Exception:
                log.exception("ocr_page_render_failed", extra={"page": index})
                results[index] = None
                continue

            text = await transcribe_image(image_b64)
            results[index] = text
            log.info("ocr_page_done", extra={"page": index, "found_text": bool(text)})
    finally:
        pdf.close()

    return results


def _extract_docx(data: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"Could not read the Word document: {exc}") from exc

    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    # Spec sheets are usually tables; skipping them would drop the useful half.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    combined = "\n".join(parts).strip()
    if not combined:
        raise ExtractionError("The document appears to be empty.")
    return combined


# The product-profile template (data/product_profile_template.md), as Pydantic
# models — this is the single source of truth for the shape; the tool-call
# JSON schema OpenAI sees is generated FROM these models (_pydantic_tool_schema),
# and the LLM's response is parsed straight back into one (model_validate_json)
# rather than hand-rolled dict/isinstance checks. That gets us three things
# for free that the old raw-dict-schema version didn't have: a single place
# that defines "what a profile looks like" (used for the schema, the parse,
# and static typing everywhere else in this module), automatic validation
# (a malformed field is a clean ValidationError, not a silent KeyError three
# functions later), and IDE/type-checker support on every profile field.
#
# Every section is optional and defaults to None — the model can genuinely
# omit what a source document doesn't cover rather than being forced to
# invent something to fill every field. `None` here means "not in this
# document", which is exactly the signal format_profile_markdown needs to
# leave that section out rather than emit an empty heading.
class ProductVariantProfile(BaseModel):
    """One product/model's full rich profile (data/product_profile_template.md's
    13-section shape) plus which specific model it describes."""

    model_name: str = Field(
        description=(
            "This variant's own name/model code as the document gives it, e.g. "
            "'Sokkia FX-201 Total Station' — not the generic series name."
        )
    )
    what_it_does: str | None = Field(
        default=None,
        description=(
            "What it does. Use null if the document genuinely does not cover this — never "
            "invent content to fill a section. Include any application/use-case write-ups "
            "the document gives (e.g. 'Boundary and Cadastral Survey', 'Topographic Survey' "
            "sections explaining what the product is used for and how) — not just a one-line "
            "functional summary. If the document spends a paragraph explaining a use case, "
            "that content belongs here, in full."
        ),
    )
    who_should_buy: str | None = Field(
        default=None,
        description=(
            "Who should buy it. Use null if the document genuinely does not cover this — "
            "never invent content to fill a section."
        ),
    )
    who_should_not_buy: str | None = Field(
        default=None,
        description=(
            "Who should NOT buy it. Use null if the document genuinely does not cover this — "
            "never invent content to fill a section."
        ),
    )
    features: str | None = Field(
        default=None,
        description=(
            "Features. Use null if the document genuinely does not cover this — never invent "
            "content to fill a section. List specs with their exact numbers/units as given in "
            "the document (ranges, accuracy figures, weights, battery life, included "
            "accessories, connectivity specs) — do not summarize a number into a generic "
            "description. 'Reflectorless range: 0.3 to 800m' stays exactly that, never becomes "
            "'reflectorless laser measurement'. If the document gives you ten specs, keep ten "
            "specs. Also include named onboard software/feature packages (e.g. an onboard "
            "field-software suite and what it does), standard package contents/included "
            "accessories, and any named feature systems (e.g. a guide-light or target-locating "
            "system) with their actual described behavior — not just their name."
        ),
    )
    benefits: str | None = Field(
        default=None,
        description=(
            "Benefits. Use null if the document genuinely does not cover this — never invent "
            "content to fill a section. Include what named features/use-cases actually mean "
            "for the buyer, drawing on any feature-explanation or comparison content in the "
            "document (e.g. a comparison against a previous/other model's accuracy or range) — "
            "do not compress a paragraph explaining a benefit into a single generic sentence "
            "if the document gives more detail than that."
        ),
    )
    price: str | None = Field(
        default=None,
        description=(
            "Price. Use null if the document genuinely does not cover this — never invent "
            "content to fill a section."
        ),
    )
    competitors: str | None = Field(
        default=None,
        description=(
            "Competitors. Use null if the document genuinely does not cover this — never "
            "invent content to fill a section."
        ),
    )
    advantages: str | None = Field(
        default=None,
        description=(
            "Advantages. Use null if the document genuinely does not cover this — never "
            "invent content to fill a section."
        ),
    )
    limitations: str | None = Field(
        default=None,
        description=(
            "Limitations. Use null if the document genuinely does not cover this — never "
            "invent content to fill a section."
        ),
    )
    common_objections: str | None = Field(
        default=None,
        description=(
            "Common objections. Use null if the document genuinely does not cover this — "
            "never invent content to fill a section."
        ),
    )
    responses: str | None = Field(
        default=None,
        description=(
            "Responses. Use null if the document genuinely does not cover this — never invent "
            "content to fill a section."
        ),
    )
    faqs: str | None = Field(
        default=None,
        description=(
            "Frequently asked questions. Use null if the document genuinely does not cover "
            "this — never invent content to fill a section."
        ),
    )
    upselling_opportunities: str | None = Field(
        default=None,
        description=(
            "Upselling opportunities. Use null if the document genuinely does not cover this "
            "— never invent content to fill a section."
        ),
    )

    def section_items(self) -> list[tuple[str, str]]:
        """(label, text) for every section that's actually filled — the
        iteration format_profile_markdown and the enrichment prompt both
        want, without either needing to know the field list themselves."""
        return [
            (label, value)
            for key, label in PROFILE_SECTION_LABELS
            if (value := getattr(self, key)) is not None
        ]

    def filled_count(self) -> int:
        return len(self.section_items())


class ProductProfileResult(BaseModel):
    """The full tool-call payload: one or more variant profiles from a single
    uploaded document. Most documents describe one model and come back as a
    single-item list; a document that genuinely gives separate models their
    own specs comes back with more than one entry — see ProductVariantProfile
    and structure_product_profile's docstring."""

    variants: list[ProductVariantProfile] = Field(
        description=(
            "One entry per distinct model/variant the document actually describes — most "
            "documents cover exactly one, so this will usually be a single-item array. Only "
            "split into multiple entries when the document genuinely gives separate specs for "
            "separate model numbers (e.g. a 'series' brochure with a spec table listing "
            "different accuracy/range/price per model). Do not split a document that just "
            "mentions related products in passing — only when it gives each one its own real "
            "specs. Never add a separate entry for the bare series name itself (e.g. no extra "
            "'FX-200' entry alongside real 'FX-201' and 'FX-202' entries) — the series name is "
            "not itself a sellable model."
        )
    )


# (field_name, display_label) for every section in ProductVariantProfile,
# excluding model_name — the one place that pairing is spelled out, reused
# by section_items(), format_profile_markdown, and the enrichment prompt so
# none of them can drift out of sync with the model itself.
PROFILE_SECTION_LABELS: list[tuple[str, str]] = [
    ("what_it_does", "What it does"),
    ("who_should_buy", "Who should buy it"),
    ("who_should_not_buy", "Who should NOT buy it"),
    ("features", "Features"),
    ("benefits", "Benefits"),
    ("price", "Price"),
    ("competitors", "Competitors"),
    ("advantages", "Advantages"),
    ("limitations", "Limitations"),
    ("common_objections", "Common objections"),
    ("responses", "Responses"),
    ("faqs", "Frequently asked questions"),
    ("upselling_opportunities", "Upselling opportunities"),
]


def _pydantic_tool_schema(model: type[BaseModel], name: str, description: str) -> dict[str, Any]:
    """An OpenAI tool-call schema generated FROM a Pydantic model's own
    json schema, rather than hand-written — keeps the schema the LLM sees
    and the model used to parse its answer permanently in sync, since
    they're now literally the same source. `ref_template`/inlining nested
    $defs isn't needed here since these models are one flat level deep, but
    model_json_schema()'s `additionalProperties` default (unset) is
    explicitly pinned to False, which OpenAI's function-calling schema
    expects for strict validation."""
    schema = model.model_json_schema()
    schema["additionalProperties"] = False
    return {
        "type": "function",
        "function": {"name": name, "description": description, "parameters": schema},
    }


PROFILE_TOOL = _pydantic_tool_schema(
    ProductProfileResult,
    "record_product_profile",
    "Record one product profile per distinct model the document describes.",
)

PROFILE_PROMPT = """You turn a raw product document (a brochure, spec sheet, or manual) into a
structured profile for a sales agent to use later.

Extract only what the document actually supports — never invent a feature, a competitor, an
objection, or an FAQ that isn't genuinely there. A section the document doesn't cover should be
left null, not padded with generic or plausible-sounding content. This matters more than
completeness: a missing section costs nothing, a fabricated one costs the sales team's
credibility the moment a customer catches it.

NEVER summarize a number, a range, or a spec into vague prose. If the source says "Reflectorless
range: 0.3 to 800m" or "Accuracy: 2 inch" or "Battery BDC72, approx. 20 hours", that exact figure
must appear in your output — not "long measurement range" or "good battery life". A sales agent
that can't quote the number a customer asked about is worse than one with no answer at all,
because it looks like it's dodging the question. Specs are the single most valuable thing a
brochure gives you; do not compress them away in the name of a tidy summary. When a document is
dense with specs (a full spec sheet), your Features section should be long and specific, not
short and generic — length is not a problem here, vagueness is.

Where the document gives you the raw material but not the finished shape — e.g. it lists specs
but never states the benefit of each one, or it never explicitly poses objections/FAQs but the
content answers ones a reasonable buyer would ask — you may synthesize a reasonable inference
FROM material that is actually present, but do not introduce facts, numbers, or claims that
aren't grounded in the document itself. If in doubt, leave the section null.

DO NOT LOSE CONTENT THAT DOESN'T NEATLY FIT ONE SECTION — a real brochure has far more than a
spec table: application/use-case write-ups ("Boundary and Cadastral Survey", "Topographic
Survey", explanations of what a specific feature does and why it helps), named software/feature
descriptions (e.g. an onboard field-software package and its capabilities), included accessories/
standard package contents, and comparison callouts against a previous/other model — all of this
is real content the sales team paid for and needs, even when it isn't phrased as a bullet-point
spec. If it doesn't fit Features, put it in Benefits (what a feature or use-case means for the
buyer) or What it does (a fuller functional description) — but it must land SOMEWHERE in your
output, not be silently dropped because it didn't fit a section's expected shape. A profile that
is shorter than what the source document actually contains is a failure, exactly the same
category of failure as inventing content that isn't there — the goal is completeness AND
accuracy, not brevity. If the source is a rich, detailed brochure, your profile should be rich
and detailed too, not compressed to a handful of one-line specs.

MULTIPLE MODELS IN ONE DOCUMENT — READ THIS CAREFULLY, IT IS A COMMON CASE, NOT AN EDGE CASE:
Before writing anything, scan the ENTIRE document text for every distinct model number/name it
mentions (e.g. "FX-201", "FX-202", "iM-55", "iM-65"). Build a mental list of every one you find.
A "series" brochure or comparison spec sheet routinely covers TWO OR MORE models in one document,
each with its own accuracy/range/price/weight numbers — e.g. a "Sokkia FX-200 series" brochure
that gives FX-201 one set of specs and FX-202 a genuinely different set. If your list has more
than one distinct model with its own specs, you MUST output one variants[] entry for EACH ONE —
this is not optional and skipping any of them is a hard failure. A response that names only the
first model and silently drops the rest is exactly as bad as inventing a fake spec: the customer
asking about the model you dropped gets no answer, or worse, gets the other model's numbers by
mistake. Re-check your own output before finishing: count the distinct model numbers in the
source text, then count your variants[] entries — if the source mentions N models, your array
must have N entries, not fewer.

Each entry's model_name must be that specific model as the document names it (e.g. "Sokkia
FX-201", never the generic series name "FX-200 Series"). Never merge two models' different
numbers into one generic profile, and never let one model's specs leak into another's entry.
Content that applies to the whole series (shared accessories, shared warranty, a general
description) may be repeated across each variant's own profile where relevant.

Do NOT add an extra variants[] entry for the series itself. "FX-200 Series" or "FX-200" is a
brand/family label, not a sellable model — if the real models are FX-201 and FX-202, your output
has exactly two entries (FX-201, FX-202), never a third one named after the family/series. Only
create an entry for something a customer could actually order by that name.

Conversely, do not manufacture variants that aren't really there — a document that just mentions
a sibling product in passing, with no separate specs of its own, describes ONE model; return a
single-item variants[] list in that ordinary case, which is most documents."""


# Same shape as rag.py's _CODE_RE/extract_codes — deliberately not imported
# from there to avoid a documents.py -> rag.py dependency for one regex;
# this is a code-*count* backstop, not a retrieval concern, so it stays
# local. Kept in sync by eye since both patterns are simple and stable.
_MODEL_CODE_RE = re.compile(r"\b([A-Z]{1,6}[-\s]?\d{1,4}[A-Z]?(?:[-–]\d[A-Z]?)?)\b")


def _model_family_prefix(machine_name: str) -> str | None:
    """The letter-prefix of the machine's own code family (e.g. "Sokkia
    FX-200 Series" -> "FX"), used to scope which codes in the raw text
    plausibly refer to a sibling model of THIS product line — a battery
    code (BDC70) or an IP rating (IP66) mentioned in the same brochure is
    not a missed variant just because it's shaped like a code too. Returns
    None if the machine name has no code-shaped token to anchor on, in
    which case the missed-variant check is skipped entirely rather than
    guessed at."""
    match = re.search(r"\b([A-Za-z]{1,6})-?\d{1,4}\b", machine_name)
    return match.group(1).upper() if match else None


def _find_model_codes(text: str, family_prefix: str | None) -> set[str]:
    """Distinct probable model codes in raw source text that share the
    product's own family prefix (see _model_family_prefix) — used only to
    sanity-check the LLM's variant count against what the document itself
    actually mentions for THIS product line, not for retrieval, and
    deliberately not for unrelated codes (battery packs, IP ratings,
    accessory part numbers) that happen to also look code-shaped.

    Also excludes the bare series/family code itself when the text has real
    numbered siblings (e.g. "FX-200" is dropped when "FX-201"/"FX-202" are
    also present) — the same "round number" heuristic as
    _drop_redundant_series_variant, so this count doesn't manufacture a
    phantom "missed variant" out of the series name and trigger a pointless
    retry when the model already correctly returned every real model.
    """
    if not family_prefix:
        return set()
    codes = {m.group(1).replace(" ", "-").upper() for m in _MODEL_CODE_RE.finditer(text)}
    family_codes = {c for c in codes if not c.isdigit() and c.split("-")[0] == family_prefix}

    def digits_of(code: str) -> str | None:
        parts = code.split("-", 1)
        return parts[1] if len(parts) == 2 and parts[1].isdigit() else None

    digit_sets = {c: digits_of(c) for c in family_codes}
    result = set()
    for code, digits in digit_sets.items():
        if digits is None:
            result.add(code)
            continue
        is_round = len(digits) >= 2 and set(digits[1:]) == {"0"}
        siblings = sum(
            1
            for other, other_digits in digit_sets.items()
            if other != code
            and other_digits is not None
            and len(other_digits) == len(digits)
            and other_digits[0] == digits[0]
            and other_digits != digits
        )
        if is_round and siblings >= 2:
            continue
        result.add(code)
    return result


def _drop_redundant_series_variant(
    profiles: list[ProductVariantProfile],
) -> list[ProductVariantProfile]:
    """Drop a bogus generic "series" entry the model sometimes adds alongside
    the real numbered variants it already extracted correctly.

    Found live on gpt-5.6-terra (not seen on gpt-4o): given a document with
    real FX-201 and FX-202 variants, it consistently added a THIRD entry
    named after the bare series ("Sokkia FX-200" / "Sokkia FX-200 Series")
    whose own content is just a restatement of "available models: FX-201
    and FX-202" — not a real distinct model, and not asked for by the
    prompt (the prompt now explicitly forbids this too, but as established
    elsewhere in this codebase a prompt instruction alone is a nudge, never
    a guarantee — this is the code-level backstop).

    Detected narrowly: a variant's own code shares the same letter prefix
    and the same number of digits as at least two OTHER variants' codes,
    AND its own digits are "round" (every digit after the first is zero —
    200, 500, not 201, 202) while the others' are not. A round number
    sharing a leading digit with two non-round siblings is the family/series
    label, not a fourth real model — e.g. "200" is dropped when "201" and
    "202" are both present, but a genuine third real model like "FX-210"
    (not round) is never touched.
    """
    if len(profiles) < 2:
        return profiles

    def parts_of(name: str) -> tuple[str, str] | None:
        match = re.search(r"([A-Za-z]{1,4})-?(\d{2,5})[A-Za-z]?", name)
        if not match:
            return None
        return match.group(1).upper(), match.group(2)

    parsed = [parts_of(p.model_name) for p in profiles]
    keep: list[ProductVariantProfile] = []
    for i, profile in enumerate(profiles):
        this = parsed[i]
        if this is None:
            keep.append(profile)
            continue
        this_prefix, this_digits = this
        # A "round" family number: every digit after the first is zero
        # (200, 100, 50 -> "500", etc.) — real model numbers in this
        # numbering style (201, 202) are never round in the same way.
        is_round = len(this_digits) >= 2 and set(this_digits[1:]) == {"0"}
        siblings = 0
        if is_round:
            for j, other in enumerate(parsed):
                if j == i or other is None:
                    continue
                other_prefix, other_digits = other
                if (
                    other_prefix == this_prefix
                    and len(other_digits) == len(this_digits)
                    and other_digits[0] == this_digits[0]
                    and other_digits != this_digits
                ):
                    siblings += 1
        if is_round and siblings >= 2:
            log.info(
                "profile_structuring_dropped_series_label",
                extra={"model_name": profile.model_name},
            )
            continue
        keep.append(profile)
    return keep or profiles


def _clean_variant(profile: ProductVariantProfile) -> ProductVariantProfile:
    """Strip whitespace and collapse an empty/whitespace-only section to
    None — a section the model returns as "" is functionally the same as
    omitting it, but would otherwise survive as a "filled" section (an
    empty string is truthy-adjacent enough to slip past a careless check)."""
    updates: dict[str, Any] = {"model_name": profile.model_name.strip()}
    for key, _ in PROFILE_SECTION_LABELS:
        value = getattr(profile, key)
        updates[key] = value.strip() if isinstance(value, str) and value.strip() else None
    return profile.model_copy(update=updates)


async def _structure_profile_call(
    raw_text: str,
    machine_name: str,
    conversation_id: str | None,
    extra_instruction: str = "",
) -> list[ProductVariantProfile] | None:
    """One structuring completion call — returns parsed variant profiles or
    None on any failure. Separated from structure_product_profile so a retry
    with a stronger nudge can reuse the same parsing/validation logic."""
    try:
        response = await complete(
            [
                {"role": "system", "content": PROFILE_PROMPT + extra_instruction},
                {
                    "role": "user",
                    "content": f"Machine: {machine_name}\n\nDocument text:\n\n{raw_text}",
                },
            ],
            tools=[PROFILE_TOOL],
            temperature=0.1,  # extraction, not conversation — low variance wins
            # A rich, detailed brochure (application write-ups, feature
            # explanations, package contents, comparison callouts — not just
            # a bare spec table) genuinely needs more room per variant than a
            # sparse spec sheet does; 4000 was found live to truncate a real
            # detailed brochure's content before it fit.
            max_output_tokens=8000,
            conversation_id=conversation_id,
        )
    except LLMUnavailableError:
        log.warning("profile_structuring_llm_unavailable", extra={"machine": machine_name})
        return None

    if not response.tool_calls:
        log.warning("profile_structuring_no_tool_call", extra={"machine": machine_name})
        return None

    try:
        result = ProductProfileResult.model_validate_json(response.tool_calls[0].function.arguments)
    except ValidationError:
        log.warning("profile_structuring_bad_json", extra={"machine": machine_name})
        return None

    profiles = [_clean_variant(v) for v in result.variants if v.model_name.strip()]
    profiles = [p for p in profiles if p.filled_count() > 0]
    if not profiles:
        log.warning("profile_structuring_no_variants", extra={"machine": machine_name})
        return None

    return _drop_redundant_series_variant(profiles) or None


async def structure_product_profile(
    raw_text: str, machine_name: str, conversation_id: str | None = None
) -> list[ProductVariantProfile] | None:
    """Turn raw extracted document text into one or more rich product-profile
    shapes (data/product_profile_template.md) via one LLM call (plus a
    conditional second call — see below).

    A single uploaded document sometimes covers more than one distinct model
    (e.g. a Sokkia "FX-200 series" brochure with separate FX-201/FX-202 specs
    in the same file) — collapsing those into one generic profile silently
    drops which spec belongs to which model, which is worse than not having
    structured the document at all. So this always returns a *list*: most
    documents genuinely describe one model and come back as a single-item
    list; only a document that actually gives separate models their own
    specs comes back with more than one entry.

    The prompt alone is not a reliable guarantee here — verified live: at
    temperature 0.1, the same two-model source text came back with only one
    variant (silently dropping the second model) in 3 of 5 runs even with an
    explicit "count the models, match the count" instruction in the prompt.
    This matches the established pattern elsewhere in this codebase (the
    catalog-dump and unrequested-price guards in app/agent.py) — a prompt
    instruction on GPT-4o is a strong nudge, never a hard guarantee, and a
    "never do X" rule that actually matters needs a code-level backstop.
    So: _find_model_codes() counts distinct model-shaped codes actually
    present in the raw source text; if the first call returned fewer
    variants than that, one retry runs with an explicit instruction naming
    exactly which codes were missed. If the retry still comes up short, the
    (partial but non-empty) result is still used — this only ever costs one
    extra completion call, and it only fires on the failure case.

    Each list entry is a ProductVariantProfile with only the sections that
    entry's source material actually supported filled in — a section the
    model left null is simply None, never a placeholder string. Returns
    None on any failure (LLM unavailable, malformed tool call, a response
    that fails Pydantic validation), so the caller can fall back to
    ingesting the raw text unchanged rather than losing the upload entirely.
    """
    profiles = await _structure_profile_call(raw_text, machine_name, conversation_id)
    if profiles is None:
        return None

    family_prefix = _model_family_prefix(machine_name)
    source_codes = _find_model_codes(raw_text, family_prefix)
    found_codes = {c for p in profiles for c in _find_model_codes(p.model_name, family_prefix)}
    missed_codes = source_codes - found_codes

    if missed_codes and len(profiles) < len(source_codes):
        log.warning(
            "profile_structuring_missed_variants",
            extra={
                "machine": machine_name,
                "variant_count": len(profiles),
                "missed_codes": sorted(missed_codes),
            },
        )
        retry_instruction = (
            "\n\nIMPORTANT — A PREVIOUS ATTEMPT AT THIS SAME DOCUMENT MISSED SOME MODELS. The "
            f"following model code(s) appear in the source text but were NOT given their own "
            f"variants[] entry last time: {', '.join(sorted(missed_codes))}. Find each of these "
            "in the document and give every one its own complete variants[] entry with its own "
            "specs, in addition to any other models you find. Do not drop any of them again."
        )
        retried = await _structure_profile_call(
            raw_text, machine_name, conversation_id, extra_instruction=retry_instruction
        )
        if retried and len(retried) > len(profiles):
            profiles = retried

    log.info(
        "profile_structured",
        extra={
            "machine": machine_name,
            "variant_count": len(profiles),
            "sections_filled": sum(p.filled_count() for p in profiles),
        },
    )
    profiles = await _enrich_missing_sections(profiles, machine_name, conversation_id)
    return profiles or None


# Sections worth enriching with general industry knowledge when the source
# document is silent on them — a spec sheet routinely has no Competitors,
# Objections, or FAQs section at all, but a sales rep still benefits from a
# generic, clearly-labeled starting point rather than a blank one. Deliberately
# excludes what_it_does/who_should_(not)_buy/features/benefits: those must
# stay grounded in the actual source document (verbatim specs matter most
# there) and are already usually filled when the document has real content —
# enriching them risks quietly overwriting a genuine "the document doesn't
# say" with a plausible-sounding guess about a specific product's behavior.
_ENRICHABLE_SECTIONS = {
    "price",
    "competitors",
    "advantages",
    "limitations",
    "common_objections",
    "responses",
    "faqs",
    "upselling_opportunities",
}

# Prepended to every enriched section so the sales team (and this codebase's
# own PATCH-to-edit flow) can tell at a glance which content came from the
# source document versus the model's own general knowledge — the client
# asked for exactly this generation, but "asked for it" is not the same as
# "should look brochure-verified when it isn't."
_AI_ESTIMATE_PREFIX = "[AI estimate — not from the source document; verify before relying on it]\n"

_ENRICHMENT_PROMPT = """You are helping a sales team fill in gaps in a product profile for
{machine_name}, a {category} product. The source brochure/spec sheet only covered some
sections — the ones listed below are genuinely missing from it.

Using your own general knowledge of this PRODUCT CATEGORY and typical products like it (NOT
by searching the web, and NOT by inventing specifics about THIS exact model that you don't
actually know), give a reasonable, clearly-general starting point for each missing section.
Write at the level of "what's typically true for this class of product," not as if you have
verified facts about this specific model — e.g. for Price, a plausible market range for this
category is fine; a specific rupee figure claimed as this exact model's price is not, since you
don't actually know it. For Competitors, name real, genuinely known competing brands/products in
this category if you know them — do not invent brand names. For Objections/FAQs, use realistic,
commonly-seen patterns for this category of purchase (budget concerns, comparison requests,
timeline questions) rather than anything specific to this model that you can't actually verify.

If you have no genuine general knowledge to offer for a section, leave it null — a missing
section is fine, a made-up one is not.

Known context about this specific product (from its actual source document, use this to make
your answer relevant, but do not contradict it):
{known_context}

Missing sections to fill: {missing_keys}"""


_SECTION_LABEL_BY_KEY = dict(PROFILE_SECTION_LABELS)


def _enrichment_model(missing_keys: list[str]) -> type[BaseModel]:
    """A Pydantic model with exactly the missing sections as fields, built on
    the fly via pydantic.create_model — the LLM is only ever asked for what's
    actually absent, never re-asked to (or able to) overwrite a section the
    source document already filled. Every field is required-but-nullable
    (str | None with no default) so the model must make an explicit null/
    string choice for each one rather than silently omitting a key."""
    fields = {
        key: (
            str | None,
            Field(
                description=(
                    f"{_SECTION_LABEL_BY_KEY[key]} — general category knowledge only, "
                    "null if you don't have any genuine basis for one."
                )
            ),
        )
        for key in missing_keys
    }
    return create_model("SectionEnrichment", **fields)  # type: ignore[call-overload]


async def _enrich_missing_sections(
    profiles: list[ProductVariantProfile], machine_name: str, conversation_id: str | None
) -> list[ProductVariantProfile]:
    """Fill in sections a source document didn't cover using the model's own
    general category knowledge — explicitly labeled as such, never presented
    as brochure-verified fact. Only touches _ENRICHABLE_SECTIONS, and only
    when at least one of them is actually missing; a profile with everything
    already filled costs nothing extra. Failures here are non-fatal — the
    caller already has a perfectly good (if sparser) profile without this."""
    enriched: list[ProductVariantProfile] = []
    for profile in profiles:
        missing = [k for k in _ENRICHABLE_SECTIONS if getattr(profile, k) is None]
        if not missing:
            enriched.append(profile)
            continue

        known_context = "\n".join(f"{label}: {text}" for label, text in profile.section_items())
        variant_label = profile.model_name or machine_name
        enrichment_model = _enrichment_model(missing)
        try:
            response = await complete(
                [
                    {
                        "role": "system",
                        "content": _ENRICHMENT_PROMPT.format(
                            machine_name=variant_label,
                            category="industrial/construction equipment",
                            known_context=known_context or "(nothing else known)",
                            missing_keys=", ".join(missing),
                        ),
                    },
                ],
                tools=[
                    _pydantic_tool_schema(
                        enrichment_model,
                        "record_enrichment",
                        "Record general-knowledge estimates for the missing sections only.",
                    )
                ],
                temperature=0.2,
                # Up to 8 enrichable sections can be missing at once (a bare
                # spec sheet with nothing beyond specs), each wanting a real
                # paragraph — found live that 1200 tokens truncated the JSON
                # mid-string on exactly that case (8/8 missing), which
                # produced a clean ValidationError ("EOF while parsing a
                # string") that was being silently swallowed below, so the
                # enrichment looked like it ran but nothing was ever filled.
                max_output_tokens=1200 * max(1, len(missing) // 2),
                conversation_id=conversation_id,
            )
        except LLMUnavailableError:
            log.warning("profile_enrichment_llm_unavailable", extra={"machine": variant_label})
            enriched.append(profile)
            continue

        if not response.tool_calls:
            enriched.append(profile)
            continue
        try:
            result = enrichment_model.model_validate_json(response.tool_calls[0].function.arguments)
        except ValidationError as exc:
            # Not just non-fatal noise: a truncated/malformed tool-call
            # response here previously vanished with zero trace (see the
            # max_output_tokens comment above) — log it so a real recurrence
            # is visible instead of silently looking like "nothing to enrich".
            log.warning(
                "profile_enrichment_bad_json",
                extra={
                    "machine": variant_label,
                    "missing": len(missing),
                    "completion_tokens": response.completion_tokens,
                    "error": str(exc)[:200],
                },
            )
            enriched.append(profile)
            continue

        updates: dict[str, Any] = {}
        for key in missing:
            value = getattr(result, key)
            if isinstance(value, str) and value.strip():
                updates[key] = _AI_ESTIMATE_PREFIX + value.strip()
        log.info(
            "profile_enriched",
            extra={"machine": variant_label, "missing": len(missing), "filled": len(updates)},
        )
        enriched.append(profile.model_copy(update=updates) if updates else profile)
    return enriched


def format_profile_markdown(
    profile: ProductVariantProfile | list[ProductVariantProfile], machine_name: str
) -> str:
    """Render structured profile(s) as `##`/`###` markdown — the same shape
    data/product_profile_template.md uses, so chunk_text's paragraph/
    heading-based splitting handles it with no changes.

    Accepts either a single variant's profile (one machine, one profile —
    the ordinary case) or a list of variant profiles for a SINGLE machine
    that genuinely has multiple types (e.g. "Sokkia FX-200 Series" covering
    FX-201 and FX-202) — the client's own mental model is one machine with
    variants underneath it, not several unrelated machines, so multiple
    variants render as `### Type: {model_name}` sub-sections nested under
    one `## {machine_name}` heading rather than as separate documents.
    """
    if isinstance(profile, ProductVariantProfile):
        return _format_single_profile_markdown(profile, heading_level="##")

    lines = [f"## {machine_name}", ""]
    for variant in profile:
        lines.append(f"### Type: {variant.model_name}")
        lines.append("")
        lines.append(_format_single_profile_markdown(variant, heading_level="####"))
        lines.append("")
    return "\n".join(lines).strip()


def _format_single_profile_markdown(profile: ProductVariantProfile, *, heading_level: str) -> str:
    """One profile's sections at the given heading depth — the shared
    renderer format_profile_markdown uses for both the single-variant case
    (### sections under a ## title) and each variant nested inside a
    multi-variant machine document (#### sections under a ### Type: title)."""
    sub_level = heading_level + "#"
    lines = []
    for label, text in profile.section_items():
        lines.append(f"{sub_level} {label}")
        lines.append(text)
        lines.append("")
    return "\n".join(lines).strip()


def chunk_text(text: str, machine_name: str) -> list[str]:
    """Split on paragraph boundaries, packing up to CHUNK_CHARS.

    Every chunk is prefixed with the machine name so a retrieved fragment is
    self-describing — otherwise a chunk of bare specifications gives the model
    no way to know which machine it belongs to.
    """
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    buffer = ""

    for paragraph in paragraphs:
        if len(buffer) + len(paragraph) + 2 <= CHUNK_CHARS:
            buffer = f"{buffer}\n\n{paragraph}" if buffer else paragraph
            continue
        if buffer:
            chunks.append(buffer)
        # A single paragraph larger than the budget gets hard-split.
        while len(paragraph) > CHUNK_CHARS:
            chunks.append(paragraph[:CHUNK_CHARS])
            paragraph = paragraph[CHUNK_CHARS - CHUNK_OVERLAP :]
        buffer = paragraph
    if buffer:
        chunks.append(buffer)

    return [f"## {machine_name}\n{chunk}" for chunk in chunks]


def _stable_point_id(key: str) -> int:
    """A Qdrant point id that is stable across processes and deploys.

    The previous version used Python's builtin `hash()`, which is
    process-randomized for strings by default (PYTHONHASHSEED) — the same
    machine+chunk-index produced a DIFFERENT id every time the backend
    restarted (a new Render deploy, a worker process cycling). The
    "re-uploading the same document updates in place" comment this scheme's
    docstring made was therefore false across restarts: a re-upload or edit
    after a deploy silently added a second, orphaned copy of every chunk
    instead of replacing the first, so a customer's answer could come from
    stale content sitting alongside the corrected version. Found while
    reviewing the multi-variant work, not something that surfaced through a
    customer report — fixed before it did. `hashlib.md5` is deterministic
    for the same input in any process, any process, forever.
    """
    import hashlib

    digest = hashlib.md5(key.encode("utf-8")).hexdigest()
    return int(digest[:16], 16) % (2**63)


async def ingest_document(
    *,
    machine_name: str,
    category: str,
    text: str,
    machine_code: str | None = None,
    machine_id: str | None = None,
    doc_type: DocumentType = DocumentType.BROCHURE,
    price_range: str | None = None,
    source_filename: str | None = None,
) -> dict[str, Any]:
    """Chunk, embed and upsert a document's text into Qdrant.

    Returns a summary the dashboard can show the user. Raises nothing on
    partial failure — the caller gets counts and can report honestly.
    """
    from qdrant_client import models

    chunks = chunk_text(text, machine_name)
    if not chunks:
        return {"chunks": 0, "embedded": 0, "error": "no text to index"}

    if not await ensure_collection():
        return {"chunks": len(chunks), "embedded": 0, "error": "vector store unavailable"}

    client = get_client()
    if client is None:
        return {"chunks": len(chunks), "embedded": 0, "error": "vector store unavailable"}

    # Strip AI-estimate paragraphs (see _AI_ESTIMATE_PREFIX) before pulling
    # codes — found live: an enriched Competitors section naming real rival
    # products ("Leica TS16/TS20", "Trimble S9") got its OWN model codes
    # picked up into THIS machine's `codes` list, so a customer asking about
    # "TS16" (a Leica product, never ours) would incorrectly exact-match this
    # machine's chunk. The stored/RAG text keeps the full enrichment content
    # either way — only the code-extraction input is narrowed. Checked with
    # `in`, not `.startswith()`: format_profile_markdown puts a "#### Label"
    # heading on its own line immediately before the section's text, so the
    # marker is never the first thing in the \n\n-joined paragraph — a
    # startswith check silently missed every real case and this codes list
    # kept leaking competitor codes despite the filter appearing to exist.
    codes_source = "\n\n".join(
        para for para in text.split("\n\n") if _AI_ESTIMATE_PREFIX.strip() not in para
    )
    codes = extract_codes(codes_source)
    if machine_code:
        codes = sorted({machine_code.upper(), *codes})

    # A re-ingest (edit, re-upload) must replace this machine's OLD chunks,
    # not just add new ones alongside them — the old scheme's point ids
    # weren't stable enough for upsert-in-place to reliably do that across a
    # process restart (see _stable_point_id). Deleting by machine_id first
    # makes this correct regardless: fewer new chunks than before leaves no
    # orphaned tail, and it degrades safely — a delete failure here just
    # means the upsert below adds alongside whatever's left, no worse than
    # the previous behavior, never a hard failure of the upload itself.
    if machine_id:
        try:
            await client.delete(
                collection_name=settings.qdrant_collection,
                points_selector=models.FilterSelector(
                    filter=models.Filter(
                        must=[
                            models.FieldCondition(
                                key="machine_id", match=models.MatchValue(value=machine_id)
                            )
                        ]
                    )
                ),
            )
        except Exception:
            log.exception("ingest_pre_delete_failed", extra={"machine_id": machine_id})

    # id_key is scoped to machine_id when known (a real UUID, so two machines
    # can never collide even if they briefly share a display name) and falls
    # back to machine_name only for the rare caller that doesn't have an id
    # yet — still deterministic, just a narrower guarantee.
    id_key_base = machine_id or machine_name

    embedded = 0
    batch_size = 16  # small batches keep peak memory low on the free tier
    for start in range(0, len(chunks), batch_size):
        batch = chunks[start : start + batch_size]
        vectors = await embed(batch)
        if len(vectors) != len(batch):
            log.error(
                "ingest_embed_mismatch",
                extra={"expected": len(batch), "got": len(vectors)},
            )
            break
        points = [
            models.PointStruct(
                # Deterministic per machine + chunk index, and stable across
                # restarts (see _stable_point_id) — combined with the
                # pre-delete above, a re-upload/edit genuinely replaces the
                # old content rather than risking a stale duplicate sitting
                # alongside it.
                id=_stable_point_id(f"{id_key_base}:{start + offset}"),
                vector=vector,
                payload={
                    "text": chunk,
                    "title": machine_name,
                    "category": category,
                    "codes": codes,
                    "machine_id": machine_id,
                    "price_range": price_range,
                    "source": source_filename or "upload",
                    "doc_type": str(doc_type),
                },
            )
            for offset, (chunk, vector) in enumerate(zip(batch, vectors, strict=True))
        ]
        try:
            await client.upsert(collection_name=settings.qdrant_collection, points=points)
            embedded += len(points)
        except Exception:
            log.exception("ingest_upsert_failed", extra={"machine": machine_name})
            break

    log.info(
        "document_ingested",
        extra={
            "machine": machine_name,
            "chunks": len(chunks),
            "embedded": embedded,
            "codes": codes[:5],
        },
    )
    return {"chunks": len(chunks), "embedded": embedded, "codes": codes}


async def ingest_accessory(
    *,
    accessory_id: str,
    name: str,
    category: str | None,
    description: str | None,
) -> dict[str, Any]:
    """Embed one accessory/part as a single Qdrant point so the agent can
    recommend it during a conversation.

    Reuses the same chunk/embed/upsert machinery as machine documents rather
    than a parallel pipeline — accessory text is short (name + description),
    so it rarely needs more than one chunk. The `record_type` payload field
    is the only new thing: a discriminator so retrieval can eventually tell
    machines and accessories apart, if that's ever needed.
    """
    from qdrant_client import models

    text = f"{name}\n{description or ''}".strip()
    if not text:
        return {"chunks": 0, "embedded": 0, "error": "no text to index"}

    if not await ensure_collection():
        return {"chunks": 1, "embedded": 0, "error": "vector store unavailable"}

    client = get_client()
    if client is None:
        return {"chunks": 1, "embedded": 0, "error": "vector store unavailable"}

    chunk = f"## {name} (accessory/part)\n{text}"
    vectors = await embed([chunk])
    if not vectors:
        log.error("accessory_ingest_embed_failed", extra={"accessory_id": accessory_id})
        return {"chunks": 1, "embedded": 0, "error": "embedding failed"}

    point = models.PointStruct(
        # Deterministic per accessory id, so re-editing updates the same
        # point in place instead of duplicating — mirrors ingest_document's
        # deterministic id scheme for machine chunks. Uses _stable_point_id
        # (hashlib, not the builtin hash()) for the same reason documented
        # there: builtin hash() is process-randomized for strings, so this
        # accessory's insert id and delete_accessory_from_index's delete id
        # would silently stop matching after any backend restart, leaving an
        # orphaned point no delete call could ever reach again.
        id=_stable_point_id(f"accessory:{accessory_id}"),
        vector=vectors[0],
        payload={
            "text": chunk,
            "title": name,
            "category": category,
            "codes": [],
            "machine_id": None,
            "accessory_id": accessory_id,
            "price_range": None,
            "source": "manual",
            "record_type": "accessory",
        },
    )
    try:
        await client.upsert(collection_name=settings.qdrant_collection, points=[point])
    except Exception:
        log.exception("accessory_ingest_upsert_failed", extra={"accessory_id": accessory_id})
        return {"chunks": 1, "embedded": 0, "error": "upsert failed"}

    # "name" is a reserved LogRecord attribute (the logger's own name) — using
    # it as an extra key raises KeyError at log time, not at review time, so
    # this was never actually caught until an accessory insert first
    # succeeded end-to-end (earlier attempts failed on a missing table before
    # ever reaching this line). accessory_name avoids the collision.
    log.info("accessory_ingested", extra={"accessory_id": accessory_id, "accessory_name": name})
    return {"chunks": 1, "embedded": 1}


async def delete_accessory_from_index(accessory_id: str) -> None:
    """Remove an accessory's Qdrant point(s).

    Deletes by payload filter on `accessory_id` rather than by computing the
    expected point id and deleting that single id directly. This matters
    because of a real migration edge: any accessory ingested before the
    _stable_point_id fix used the old, process-randomized `hash()` scheme,
    so its actual stored point id does not equal what this function would
    compute today — a by-id delete would silently delete nothing and leave
    that accessory's stale point behind forever. A payload filter finds the
    point regardless of which id scheme created it, at the one-time cost of
    a filtered scan instead of a direct point lookup (accessories are few
    and this call is not on any hot path)."""
    from qdrant_client import models

    client = get_client()
    if client is None:
        return
    try:
        await client.delete(
            collection_name=settings.qdrant_collection,
            points_selector=models.FilterSelector(
                filter=models.Filter(
                    must=[
                        models.FieldCondition(
                            key="accessory_id", match=models.MatchValue(value=accessory_id)
                        )
                    ]
                )
            ),
        )
    except Exception:
        log.exception("accessory_index_delete_failed", extra={"accessory_id": accessory_id})


async def delete_machine_from_index(machine_id: str) -> None:
    """Remove every Qdrant chunk belonging to a machine, by payload filter
    rather than a deterministic id — unlike an accessory, a machine's
    document can chunk into many points (chunk_text splits on paragraph
    boundaries), so there's no single id to target. Best-effort: deleting
    the Postgres row must not be blocked by a Qdrant hiccup, since the
    machine is already gone from the catalog either way; a stray orphaned
    chunk is a smaller problem than losing the delete entirely."""
    from qdrant_client import models

    client = get_client()
    if client is None:
        return
    try:
        await client.delete(
            collection_name=settings.qdrant_collection,
            points_selector=models.FilterSelector(
                filter=models.Filter(
                    must=[
                        models.FieldCondition(
                            key="machine_id", match=models.MatchValue(value=machine_id)
                        )
                    ]
                )
            ),
        )
    except Exception:
        log.exception("machine_index_delete_failed", extra={"machine_id": machine_id})


async def add_machine_from_document(
    *,
    name: str,
    category: str,
    data: bytes,
    filename: str,
    content_type: str | None,
    machine_code: str | None = None,
    description: str | None = None,
    price_range: str | None = None,
    lead_time: str | None = None,
    doc_type: DocumentType = DocumentType.BROCHURE,
) -> dict[str, Any]:
    """Full pipeline: extract, restructure into the rich product-profile
    shape, persist the machine and document, index for RAG.

    The client only wants to type one machine name and upload one document —
    everything else (What it does / Who should buy it / Objections /
    Responses / FAQs / etc., data/product_profile_template.md's shape) comes
    from structure_product_profile analysing the extracted text.

    That analysis can find more than one distinct model/type in a single
    document (a "series" brochure listing separate specs per model, e.g.
    FX-201 vs FX-202) — per the client's own mental model, these are TYPES
    of one machine, not separate machines: this creates exactly ONE
    `machines` row (named after what was typed at upload, e.g. "Sokkia
    FX-200 Series"), one `machine_document`, and one RAG-ingested document,
    with each detected variant rendered as its own `### Type: {model_name}`
    sub-section inside that single document (format_profile_markdown).
    Retrieval still tells the variants apart correctly: `extract_codes` picks
    up every variant's own model code from the combined document text (both
    "FX-201" and "FX-202" appear in it), so an exact-code match on a
    customer's specific question still prioritizes the right chunk — the
    codes distinguish the *content*, they just no longer need a separate
    `machine_id` to do it.

    Falls back to storing the raw extracted text unchanged if structuring
    fails entirely, so an LLM hiccup costs richness, never the whole upload.
    """
    text = await extract_text(data, filename, content_type)
    profiles = await structure_product_profile(text, name)

    resolved_code = machine_code or name.upper().replace(" ", "-")[:40]
    machine_id = await store.upsert_machine(
        machine_code=resolved_code,
        name=name,
        category=category,
        description=description,
        price_range=price_range,
        lead_time=lead_time,
    )

    if not profiles:
        stored_text = text
        variants_detected = 1
        sections_filled = 0
    elif len(profiles) == 1:
        stored_text = format_profile_markdown(profiles[0], name)
        variants_detected = 1
        sections_filled = profiles[0].filled_count()
    else:
        stored_text = format_profile_markdown(profiles, name)
        variants_detected = len(profiles)
        sections_filled = sum(p.filled_count() for p in profiles)

    await store.save_machine_document(
        machine_id=machine_id,
        doc_type=doc_type,
        title=filename,
        content=stored_text,
    )
    result = await ingest_document(
        machine_name=name,
        category=category,
        text=stored_text,
        machine_code=resolved_code,
        machine_id=machine_id,
        doc_type=doc_type,
        price_range=price_range,
        source_filename=filename,
    )

    return {
        "machine_id": machine_id,
        "name": name,
        "characters_extracted": len(text),
        "profile_sections_filled": sections_filled,
        "variants_detected": variants_detected,
        "variant_names": (
            [p.model_name for p in profiles] if profiles and variants_detected > 1 else []
        ),
        **result,
    }
